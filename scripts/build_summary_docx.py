"""기말과제 프로젝트 요약문 .docx 생성 — python-docx 기반.
일회용 스크립트. 출력: 이용주-기말과제-프로젝트요약문.docx
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = Path("이용주-기말과제-프로젝트요약문.docx")

HEADER_BG = "D9E2F3"  # 부드러운 블루
HEADER_TEXT = "1F3864"
BORDER_COLOR = "BFBFBF"


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_table_borders(table, color: str = BORDER_COLOR, size: str = "4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tbl_pr.append(borders)


def set_run_font(run, size=10, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 0:
        set_run_font(run, size=18, bold=True, color="1F3864")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_run_font(run, size=13, bold=True, color="1F3864")
    else:
        set_run_font(run, size=11, bold=True)
    return p


def add_para(doc, text, size=10, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    run.italic = italic
    return p


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True

    # 헤더 행
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=HEADER_TEXT)
        set_cell_bg(cell, HEADER_BG)

    # 데이터 행
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(val)
            set_run_font(run, size=9)

    # 열 너비 (선택)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    set_table_borders(table)
    return table


def main():
    doc = Document()

    # 페이지 여백 줄이기 (1~2 페이지 분량 확보)
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rPr.append(rFonts)

    # 표지
    add_heading(doc, "기말과제 프로젝트 요약문 — TubeBrief", level=0)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["작성자 / 학번", "이용주 / _(제출 전 본인 학번 기입)_"],
            ["프로젝트명", "TubeBrief — 유튜브 구독 채널 자동 요약"],
            ["GitHub", "https://github.com/leeyongjoo9-rgb/tubebrief_mvp"],
            ["서비스 URL", "https://tubebrief-mvp.vercel.app/"],
        ],
        col_widths_cm=[3.5, 13.5],
    )

    # 1. 서비스 개요
    add_heading(doc, "1. 서비스 개요", level=1)
    add_para(
        doc,
        "구독한 유튜브 채널의 신규 영상을 매일 자동 감지해 AI 로 구조화된 한 장짜리 요약 보고서로 변환·아카이빙하는 개인용 콘텐츠 다이제스트.",
    )

    # 2. 문제와 사용자
    add_heading(doc, "2. 문제와 사용자", level=1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["사용자", "구독 채널 8개 이상 / 콘텐츠 중심 시청 / 시간 부족한 개인"],
            ["상황", "매일 채널마다 1~5편 신규 영상 누적. 무엇부터 볼지 판단 어려움"],
            ["기존 한계", "유튜브 자막은 길고 비구조화. 한 줄 캡션은 정보 부족"],
            ["차별점", "헤드라인·본문·출연자·거론된 인물·기업·시각 점프 링크 분리된 한 장 보고서"],
        ],
        col_widths_cm=[2.7, 14.3],
    )

    # 3. 핵심 기능
    add_heading(doc, "3. 핵심 기능", level=1)
    add_table(
        doc,
        ["기능", "설명"],
        [
            ["자동 감지 cron", "Vercel Cron 일 1회. RSS → 자막 (실패 시 메타데이터 폴백) → LLM 요약"],
            ["코너 단위 구독", "제목 + RSS 설명문 함께 키워드 AND 필터. 설명문 해시태그(#곽재식)로 호스트 매치"],
            ["출연자 ↔ 거론 인물 분리", "‘직접 출연’ / ‘대화 중 거론’ 별도 필드. 환각 방지 + 채널 호스트 자동 제외"],
        ],
        col_widths_cm=[3.5, 13.5],
    )
    add_para(doc, "비고: 검색·필터·텍스트 내보내기는 v2 후보 (미구현).", italic=True, size=9)

    # 4. 기술스택
    add_heading(doc, "4. 기술스택", level=1)
    add_table(
        doc,
        ["영역", "선택", "선택 이유"],
        [
            ["Framework", "Next.js 16 (App Router)", "API Routes 로 백엔드 통합. Vercel native"],
            ["언어", "TypeScript", "LLM 응답 스키마와 코드 타입 일치"],
            ["스타일", "Tailwind v4 + shadcn/ui", "AI 코딩 친화, 디자인 시스템 즉시 사용"],
            ["DB", "Supabase (PostgreSQL)", "jsonb 로 LLM 출력 그대로 저장. RLS 보안"],
            ["LLM", "OpenAI gpt-5-mini", "structured outputs JSON Schema strict 강제. 한국어 분량 준수"],
            ["자막", "youtube-transcript + 폴백", "클라우드 IP 봇 차단 회피용 폴백"],
            ["신규 감지", "YouTube RSS", "API 키·쿼터 불필요"],
            ["배포", "Vercel (Hobby)", "Cron + 함수 + 정적 호스팅 + GitHub 자동 빌드"],
        ],
        col_widths_cm=[2.5, 4.5, 10.0],
    )

    # 5. MCP/API 활용
    add_heading(doc, "5. MCP / API 활용", level=1)
    add_table(
        doc,
        ["API / Platform", "활용 목적", "연동 흐름"],
        [
            ["OpenAI Chat (gpt-5-mini)", "자막 → 구조화 JSON 요약", "자막 + 시스템 프롬프트 + JSON Schema → API → Supabase jsonb 저장"],
            ["Supabase REST + JS SDK", "채널/구독/영상/요약 저장 조회", "서버는 service role, 클라이언트는 publishable key (RLS)"],
            ["YouTube RSS", "신규 영상 감지", "feeds/videos.xml 폴링 → XML 파싱 → 키워드 매치 → INSERT"],
            ["YouTube 페이지 스크래핑", "URL → channel_id 추출", "6 패턴 fallback (canonical / og:url / Schema.org / JSON channelId)"],
            ["Vercel Platform (인프라)", "Cron 자동화 + Serverless Functions + 환경변수 + GitHub 자동 빌드", "vercel.json 의 crons → 매일 0시 UTC /api/cron 자동 트리거. push 시 자동 재배포 (GitHub webhook)"],
        ],
        col_widths_cm=[4.0, 4.0, 9.0],
    )
    add_para(
        doc,
        "한계: LLM 비용 영상당 ~$0.025, Vercel Hobby 60초 timeout. → 보수적 limit + temperature 0.2 로 retry 최소화.",
        italic=True,
        size=9,
    )

    # 6. 컨텍스트 엔지니어링
    add_heading(doc, "6. 컨텍스트 엔지니어링", level=1)
    add_table(
        doc,
        ["Context 파일", "주요 역할", "실제 활용 방식"],
        [
            ["CLAUDE.md", "절대 규칙 6개 + 작업 순서 8단계 + 코딩 스타일", "규칙 1번(키 보안) → 서버 액션 + service role 분리. 규칙 6번(<form> 금지) → onClick 사용"],
            ["TubeBrief_PRD.md", "문제·기능 P0/P1/P2·사용자 흐름·파이프라인", "자막 폴백 3단계 그대로 코드 구현 (transcript_ok / metadata_fallback / failed)"],
            ["SUMMARY_SCHEMA.json", "LLM 출력 JSON 스키마 진실 공급원", "OpenAI response_format.json_schema 에 그대로 전달. strict mode 강제"],
            ["PRESENTATION.md", "발표 콘텐츠 원자료 + 옛/새 규칙 비교", "9슬라이드 70% 사전 정리. 비교용 두 URL 박혀 있어 GitHub 리뷰 시 노출"],
        ],
        col_widths_cm=[3.5, 5.0, 8.5],
    )

    # 7. 개발기록
    add_heading(doc, "7. 개발기록 — 가장 어려웠던 점과 해결", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("문제: ")
    set_run_font(r1, size=10, bold=True)
    r2 = p.add_run(
        "LLM 이 한국어 요약에서 약속한 분량·문단·시각 마커 규칙을 끝까지 안 지킴. gpt-4o-mini 평균 brief 290자 / 1문단 / (mm:ss) 마커 0개 (목표 1000자+ / 5~8문단 / 2~5마커)."
    )
    set_run_font(r2, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("해결 과정: ")
    set_run_font(r1, size=10, bold=True)
    r2 = p.add_run(
        "(1) 프롬프트를 ‘가이드’ → ‘반드시 ~’ 명령조로 강화 → (2) validateSummary 후처리 검증 + 위반 시 자동 재시도 (MAX 2회) → (3) gpt-5-mini 로 모델 교체."
    )
    set_run_font(r2, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("결과: ")
    set_run_font(r1, size=10, bold=True)
    r2 = p.add_run(
        "같은 영상(인터넷 검열, 자막 13,641자)에서 brief 264자 / 1문단 / 시각마커 0 → "
    )
    set_run_font(r2, size=10)
    r3 = p.add_run("2,122자 / 10문단 / 시각마커 5개 + 기업 7개 추출")
    set_run_font(r3, size=10, bold=True)
    r4 = p.add_run(
        ". 약 8배 정보 밀도 증가. 두 영상을 라이브 비교 가능하도록 보존 (8Khu3IoZric vs muA-MB0k7SE — tubebrief-mvp.vercel.app/videos/{video_id})."
    )
    set_run_font(r4, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("배운 점: ")
    set_run_font(r1, size=10, bold=True)
    r2 = p.add_run(
        "프롬프트 단독으로는 부족. 검증·재시도·모델 선택의 3겹 가드레일이 신뢰할 수 있는 LLM 시스템의 필수 구성."
    )
    set_run_font(r2, size=10)

    # 8. 한계와 개선
    add_heading(doc, "8. 한계와 개선", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("운영 비용 / 데이터: ")
    set_run_font(r1, size=10, bold=True)
    r2 = p.add_run(
        "영상 mp4 는 저장 안 함 (YouTube 호스팅). DB 에는 영상 1편당 ~20~55KB (메타 + 자막 + 요약). "
        "Supabase Free 500MB 안에서 약 1만 편 보관 가능 → 월 100편 처리 시 8~10년 무료 운영. "
        "실제 월 운영비 ≈ $2.5 (OpenAI 만), Supabase/Vercel/GitHub 모두 무료 tier."
    )
    set_run_font(r2, size=10)

    add_table(
        doc,
        ["한계", "향후 개선"],
        [
            [
                "신규 등록 채널의 RSS 15편 이전 영상 백필 불가 (단, 일일 cron 운영 중에는 신규 영상 누락 없음 → 정상 운영에선 한계 아님)",
                "YouTube Data API 도입 시 등록 직후 백필 가능 (쿼터·인증 부담)",
            ],
            ["Vercel Hobby 60초 / 일 1회 cron", "Pro 업그레이드 또는 외부 워커 분리"],
            ["자막 음성 인식 오류 (곽재식→곽제식)", "description 해시태그를 정답 표기로 LLM 에 전달"],
            ["인증 부재 — 배포 URL 공개", "Vercel Password Protection 또는 Supabase Auth"],
            ["검색·필터·텍스트 내보내기 미구현", "v2 후보"],
            ["모바일 최적화 부족", "v0.dev 기반 재설계"],
        ],
        col_widths_cm=[8.5, 8.5],
    )

    doc.save(str(OUTPUT))
    print(f"saved: {OUTPUT.resolve()}")
    print(f"size: {OUTPUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
